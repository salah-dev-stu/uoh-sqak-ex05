"""Fill the submission .docx template with the locked metadata and export to PDF.

Final values (placeholders — orchestrator will recalibrate before submission):
  group_code = "uoh-sqak"   (Salah Qadah + Andalus Kalash — same pair as HW1+HW2)
  self_grade = 85           (HW1 self-graded 90 vs actual 85.54; HW2 not yet graded — stay calibrated)
  github     = https://github.com/salah-dev-stu/uoh-sqak-ex05  (confirm with Andalus)
  late       = no

Moodle assignment id=278465 — deadline Friday 26 June 2026, 23:59 Asia/Jerusalem

Run:  uv run python scripts/fill_submission_pdf.py
Outputs:
  uoh-sqak-ex05.docx   (filled cover sheet)
  uoh-sqak-ex05.pdf    (final; what each pair member uploads separately to Moodle)
"""

from __future__ import annotations

import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document

REPO = Path(__file__).resolve().parent.parent
TEMPLATE = REPO / "uoh-rl07-ex05-template.docx"

GROUP_CODE = "uoh-sqak"
SELF_GRADE = "85"
EXERCISE_NUMBER = "05"
GITHUB_URL = "https://github.com/salah-dev-stu/uoh-sqak-ex05"
LATE = "no"

STUDENT_1 = {
    "id": "323039974",
    "first_en": "Salah",
    "last_en": "Qadah",
    "first_he": "סלאח",
    "last_he": "קדח",
}
STUDENT_2 = {
    "id": "211435797",
    "first_en": "Andalus",
    "last_en": "Kalash",
    "first_he": "אנדלוס",
    "last_he": "כלש",
}


def append_value(paragraph, value: str) -> None:
    """Append " <value>" as a new run to a paragraph (preserves label run)."""
    paragraph.add_run(" " + value)


def main() -> int:
    if not TEMPLATE.exists():
        print(f"FAIL: template not found at {TEMPLATE}", file=sys.stderr)
        return 1
    intermediate = REPO / f"{GROUP_CODE}-ex05.docx"
    shutil.copy(TEMPLATE, intermediate)

    d = Document(intermediate)
    p = d.paragraphs

    append_value(p[0], EXERCISE_NUMBER)  # exercise number
    append_value(p[2], GROUP_CODE)  # group ID code
    append_value(p[4], SELF_GRADE)  # self-grade
    append_value(p[7], STUDENT_1["id"])  # S1 ID
    append_value(p[8], STUDENT_1["first_en"])
    append_value(p[9], STUDENT_1["last_en"])
    append_value(p[10], STUDENT_1["first_he"])
    append_value(p[11], STUDENT_1["last_he"])
    append_value(p[14], STUDENT_2["id"])  # S2 ID
    append_value(p[15], STUDENT_2["first_en"])
    append_value(p[16], STUDENT_2["last_en"])
    append_value(p[17], STUDENT_2["first_he"])
    append_value(p[18], STUDENT_2["last_he"])
    append_value(p[20], GITHUB_URL)  # GitHub link
    append_value(p[21], LATE)  # late yes/no

    d.save(intermediate)
    print(f"  filled docx: {intermediate}")

    pdf = REPO / f"{GROUP_CODE}-ex05.pdf"
    # Arial Unicode MS preserves ASCII hyphen (U+002D) in URLs and group code,
    # which Lucida Grande silently rewrites to U+2011 (non-breaking hyphen) —
    # that breaks regex-based URL extraction by an automated grader. Arial
    # Unicode MS also has full Hebrew coverage (no missing-character warnings).
    rc = subprocess.run(
        [
            "pandoc",
            str(intermediate),
            "-o",
            str(pdf),
            "--pdf-engine=xelatex",
            "-V",
            "mainfont=Arial Unicode MS",
            "-V",
            "geometry:margin=1in",
        ],
        capture_output=True,
        text=True,
    )
    if rc.returncode != 0:
        print("pandoc xelatex failed; STDERR:")
        print(rc.stderr)
        return 1
    print(f"  pdf written: {pdf}  ({pdf.stat().st_size:,} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
